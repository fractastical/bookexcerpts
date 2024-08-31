import os
import re
from docx import Document
from PIL import Image
import io

def sanitize_filename(filename, max_length=100):
    sanitized = re.sub(r'\W+', '_', filename)
    return sanitized[:max_length]

def estimate_page_number(para_index, paragraphs_per_page=30):
    return (para_index // paragraphs_per_page) + 1

def extract_and_process_images(doc_path, output_folder):
    doc = Document(doc_path)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    image_number = 1
    pending_image = None
    fallback_counter = 1

    for i, para in enumerate(doc.paragraphs):
        page_number = estimate_page_number(i)

        # Check for images in the paragraph
        for run in para.runs:
            for drawing in run._element.xpath('.//a:blip'):
                rId = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId:
                    image = doc.part.related_parts[rId]
                    
                    # Open image in PIL and convert to grayscale
                    image_bytes = io.BytesIO(image._blob)
                    img = Image.open(image_bytes)
                    img = img.convert('L')  # Convert to grayscale

                    # Determine if the image is small (under 10KB)
                    image_size = len(image_bytes.getvalue())
                    prefix = f"{page_number}_"
                    if image_size < 10240:  # 10KB = 10240 bytes
                        prefix = f"micro_{prefix}"

                    # Store the image until the next citation is found
                    pending_image = (img, prefix)
                    print("Found an image, waiting for the corresponding citation...")  # Debugging print

        # If there's a pending image, look for the next citation
        if pending_image:
            img, prefix = pending_image
            match = re.match(r'^\s*(.*?)\s*\((\d{4})\)\s*(.*?)\.?\s*$', para.text.strip())
            if match:
                author = match.group(1).strip()
                date = match.group(2).strip()
                title = match.group(3).strip()
                citation = f"{author} ({date}) {title}"
                citation = sanitize_filename(citation)
                print(f"Found citation for image: {citation}")

                # Save the image with the citation name
                image_name = f"{prefix}{citation}.png"
                img.save(os.path.join(output_folder, image_name), format='PNG')
                print(f"Saved image {image_name}")

                # Reset pending image after saving
                pending_image = None
            else:
                # Fallback to using the first few words of the next paragraph as the image title
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1].text.strip()
                    if next_para:
                        fallback_title = " ".join(next_para.split()[:5])
                        fallback_title = sanitize_filename(fallback_title)
                        image_name = f"{prefix}{fallback_title}.png"
                        print(f"Using fallback title for image: {image_name}")

                        # Save the image with the fallback title
                        img.save(os.path.join(output_folder, image_name), format='PNG')
                        print(f"Saved image {image_name}")

                        # Reset pending image and increase fallback counter
                        pending_image = None
                        fallback_counter += 1

# Define the path to the document and output folder
doc_path = 'cyberutopias-r.docx'
output_folder = 'images'

# Extract and process images
extract_and_process_images(doc_path, output_folder)
