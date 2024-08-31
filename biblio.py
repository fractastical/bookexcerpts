import re
import json
from docx import Document
from collections import defaultdict
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def format_citation_mit(citation):
    author = citation['author']
    date = citation['date']
    title = citation['title']
    
    # Format author names
    authors = author.split(' and ')
    if len(authors) > 1:
        formatted_authors = ', '.join(authors[:-1]) + ' and ' + authors[-1]
    else:
        formatted_authors = authors[0]
    
    # Ensure correct capitalization
    def capitalize_title(title):
        words = title.split()
        result = []
        for i, word in enumerate(words):
            if i == 0 or word[0] in ['"', "'"]:  # Capitalize first word or words after quotes
                result.append(word.capitalize())
            elif len(word) > 3 and word.lower() not in ['a', 'an', 'the', 'and', 'but', 'or', 'for', 'nor', 'on', 'at', 'to', 'from', 'by', 'in', 'of']:
                result.append(word.capitalize())
            else:
                result.append(word.lower())
        return ' '.join(result)

    capitalized_title = capitalize_title(title)

    return f"{formatted_authors}. {date}. {capitalized_title}."

def extract_citations(doc):
    main_text_pattern = r'^\s*(.*?)\s*\((\d{4}(?:-\d{4})?)\)\.?\s*(.*?)(?:,\s*p\.\s*(\d+))?\.?\s*$'
    footnote_patterns = [
        r'(\w+),\s+(.*?)\s*\((\d{4}(?:-\d{4})?)\)\.\s*(.*?)\.',  # Author, Title (Year). Publisher.
        r'(\w+)\s+(\w+)\s+(?:&|and)\s+(\w+)\s+(\w+)\s*\((\d{4}(?:-\d{4})?)\)\.\s*(.*?)\.',  # Author1 Lastname1 & Author2 Lastname2 (Year). Title.
        r'(\w+),\s+(.*?)\s*(\d{4}(?:-\d{4})?)\.',  # Author, Title Year.
        r'(\w+),\s+(\w+)\s*\((\d{4}(?:-\d{4})?)\)\s*(.*?)\.',  # Lastname, Firstname (Year) Title.
    ]
    bibliography_pattern = r'(.*?)\s*\((\d{4}(?:-\d{4})?)\)\.?\s*(.*?)\.\s*'
    
    citations = defaultdict(lambda: {'main_text': [], 'footnotes': [], 'extracts': [], 'references': []})
    chapter_counter = 0  # Start with Chapter 1
    headings = {
        "chapter": "",
        "heading2": "",
        "heading3": "",
        "heading4": ""
    }

    in_bibliography = False
    bibliography_section = ""
    bibliography_headings = {"extracts", "references", "additional readings"}

    def process_footnote(footnote_text):
        standard_format = False
        for pattern in footnote_patterns:
            match = re.search(pattern, footnote_text)
            if match:
                groups = match.groups()
                if len(groups) == 4:  # First pattern
                    author, title, date, publisher = groups
                elif len(groups) == 6:  # Second pattern
                    author = f"{groups[0]} {groups[1]} and {groups[2]} {groups[3]}"
                    date, title = groups[4], groups[5]
                elif len(groups) == 3:  # Third pattern
                    author, title, date = groups
                elif len(groups) == 4:  # Fourth pattern
                    author = f"{groups[1]} {groups[0]}"
                    date, title = groups[2], groups[3]
                
                citation_info = {
                    "author": author,
                    "date": date,
                    "title": title,
                    "standard_format": True
                }
                standard_format = True
                break
        
        if not standard_format:
            citation_info = {
                "full_text": footnote_text,
                "standard_format": False
            }
        
        return citation_info

    # Process footnotes
    footnotes = doc.part._element.findall('.//w:footnote', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    for footnote in footnotes:
        if footnote.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id'] not in ['0', '-1']:  # Skip separator and continuation separator footnotes
            footnote_text = ''.join(paragraph.text for paragraph in footnote.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
            citation_info = process_footnote(footnote_text)
            citations[headings['chapter']]['footnotes'].append(citation_info)

    for para in doc.paragraphs:
        if para.style.name == 'Heading 2':
            match = re.match(r'(\d+)\.\s*(.*)', para.text.strip())
            if match:
                chapter_counter += 1
                chapter_title = match.group(2)
            else:
                chapter_title = para.text.strip()
            
            headings['chapter'] = f"Chapter {chapter_counter}: {chapter_title}"
            headings['heading2'] = headings['chapter']
            headings['heading3'] = ""
            headings['heading4'] = ""
            in_bibliography = False

        elif para.style.name in ['Heading 3', 'Heading 4']:
            lower_text = para.text.strip().lower()
            if lower_text in bibliography_headings:
                in_bibliography = True
                bibliography_section = para.text.strip().lower()
            else:
                headings[para.style.name.lower()] = para.text.strip()
                if not in_bibliography:
                    in_bibliography = False

        if in_bibliography and para.text.strip() and para.style.name not in ['Heading 3', 'Heading 4']:
            bib_match = re.search(bibliography_pattern, para.text.strip())
            if bib_match:
                author = bib_match.group(1).strip()
                date = bib_match.group(2).strip()
                title = bib_match.group(3).strip()
                entry = {
                    "author": author,
                    "date": date,
                    "title": title
                }
                if bibliography_section == 'extracts':
                    citations[headings['chapter']]['extracts'].append(entry)
                elif bibliography_section == 'references':
                    citations[headings['chapter']]['references'].append(entry)
        elif not in_bibliography:
            match = re.match(main_text_pattern, para.text.strip())
            if match:
                author = match.group(1).strip()
                date = match.group(2).strip()
                title = match.group(3).strip()
                page_number = match.group(4)

                citation_info = {
                    "author": author,
                    "date": date,
                    "title": title,
                    "page_number": page_number if page_number else None
                }

                citations[headings['chapter']]['main_text'].append(citation_info)

    return citations

def generate_chapter_based_citations(citations):
    new_doc = Document()

    # Add a title to the document
    title = new_doc.add_paragraph("Combined Citations by Chapter")
    title_style = new_doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title.style = title_style
    title.style.font.size = Pt(18)
    title.style.font.bold = True

    # Create a single citation style
    citation_style = new_doc.styles.add_style('CitationStyle', WD_STYLE_TYPE.PARAGRAPH)
    citation_style.font.size = Pt(12)
    citation_style.font.name = 'Times New Roman'
    citation_style.paragraph_format.left_indent = Pt(36)
    citation_style.paragraph_format.first_line_indent = Pt(-36)
    citation_style.paragraph_format.space_after = Pt(12)

    for chapter, chapter_data in citations.items():
        # Add chapter heading
        chapter_heading = new_doc.add_paragraph(chapter)
        if f'ChapterHeading_{chapter}' not in new_doc.styles:
            heading_style = new_doc.styles.add_style(f'ChapterHeading_{chapter}', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
            heading_style.font.name = 'Times New Roman'
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        chapter_heading.style = new_doc.styles[f'ChapterHeading_{chapter}']

        chapter_citations = []
        
        # Add citations from extracts
        chapter_citations.extend(chapter_data['extracts'])
        
        # Add citations from main text
        for citation in chapter_data['main_text']:
            if citation not in chapter_citations:
                chapter_citations.append(citation)
        
        # Add citations from footnotes
        for footnote in chapter_data['footnotes']:
            if footnote.get('standard_format', False) and footnote not in chapter_citations:
                chapter_citations.append(footnote)

        # Sort citations alphabetically by author
        sorted_citations = sorted(chapter_citations, key=lambda x: x['author'].split()[-1])

        # Add each citation in MIT format
        for citation in sorted_citations:
            citation_text = format_citation_mit(citation)
            paragraph = new_doc.add_paragraph(citation_text)
            paragraph.style = citation_style

    # Save the new document
    new_doc.save('combined_citations_by_chapter.docx')
    print("Saved combined_citations_by_chapter.docx")

# Load the .docx file
doc_path = 'cyberutopias-r.docx'
doc = Document(doc_path)

if not doc:
    print("Failed to load the document.")
else:
    print("Document loaded successfully.")

citations = extract_citations(doc)

# Generate the chapter-based citation document
generate_chapter_based_citations(citations)
